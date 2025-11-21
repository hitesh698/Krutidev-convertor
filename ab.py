# ab.py -- Streamlit-friendly version of your converter
# (Original GUI removed. Conversion logic preserved.)
# Exposes:
#   - convert(text: str) -> str       : returns converted Kruti output (string)
#   - build_docx_bytes(text: str) -> bytes : returns a .docx file as bytes (for download)

import re
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import io

# ---------------------------------------------------------
#      UNIVERSAL SPACE NORMALIZER
# ---------------------------------------------------------
def normalize_spaces(text):
    all_spaces = [
        "\u0009", "\u000A", "\u000B", "\u000C", "\u000D",
        "\u0085", "\u00A0", "\u1680", "\u180E", "\u2000",
        "\u2001", "\u2002", "\u2003", "\u2004", "\u2005",
        "\u2006", "\u2007", "\u2008", "\u2009", "\u200A",
        "\u200B", "\u2028", "\u2029", "\u202F", "\u205F",
        "\u3000"
    ]
    for s in all_spaces:
        text = text.replace(s, " ")
    return text


# ---------------------------------------------------------
#                   KRUTI DEV BASE MAP
# ---------------------------------------------------------
base_map = {
    "अ": "v", "आ": "vk", "इ": "b", "ई": "bZ",
    "उ": "m", "ऊ": "Å", "ए": ",", "ऐ": ",s",
    "ओ": "vks", "औ": "vkS",

    "क": "d", "क्": "D",
    "ख": "[k", "ख्": "[",
    "ग": "x", "ग्": "X",
    "घ": "?k", "घ्": "?",

    "च": "p", "च्": "P",
    "छ": "N", "छ्": "n",
    "ज": "t", "ज्": "T",
    "झ": "H", "झ्": "h",

    "ट": "V", "ट्": "v",
    "ठ": "B", "ठ्": "b",
    "ड": "M", "ड्": "m",
    "ढ": "<", "ढ्": "<",
    "ण": ".k", "ण्": ".", 

    "त": "r", "त्": "R",
    "थ": "Fk", "थ्": "f",
    "द": "n", "द्": "N",
    "ध": "/k", "ध्": "/",
    "न": "u", "न्": "U",

    "प": "i", "प्": "I",
    "फ": "Q", "फ्": "q",
    "ब": "c", "ब्": "C",
    "भ": "Hk", "भ्": "H",
    "म": "e", "म्": "E",

    "य": ";", "य्": ":",
    "र": "j", "र्": "J",
    "ल": "y", "ल्": "Y",
    "व": "o", "व्": "O",

   "श": "'k", "श्": "'",
   "ष": "\"k", "ष्": "\"",
    "स": "l", "स्": "L",
    "ह": "g", "ह्": "G",

    "क्ष": "†",
    "त्र": "=",
    "ज्ञ": "?k",
    "श्र": "J",

    "ं": "a", "ः": ":", "ँ": "˜",
}


matras = {
    "ा": "k", "ि": "f", "ी": "h",
    "ु": "q", "ू": "w",
    "े": "s", "ै": "S",
    "ो": "ks", "ौ": "kS",
    "ृ": "~"
}


# ---------------------------------------------------------
#                HINDI → KRUTI DEV
# ---------------------------------------------------------
def convert_hindi(txt):
    # ----- SPECIAL FIX ONLY FOR "त्रि" -----
    txt = txt.replace("त्रि", "f=")

    # ---------------------------
    # SPECIAL RULES
    # ---------------------------
    special_map = {
        "श्र": "J",
        "र्या": ";kZ",
        "त्रि": "f=",
        "त्र": "=",
        "द्व": "}",
        "रू": ":",
        "प्र": "iz",
        "क्र": "dz",
        "द्र": "nz",
        "कृ": "d~",
        "तः": "r%",
        "त्ति": "fRr",
    }

    for u, k in special_map.items():
        txt = txt.replace(u, k)

    # ----------------------------------------------------
    # HALF-RA RULE:  र् + X  →  XZ
    # ----------------------------------------------------
    def half_ra(match):
        cons = match.group(1)
        if cons in base_map:
            return base_map[cons] + "Z"
        return cons

    txt = re.sub(r"र्([क-ह])", half_ra, txt)

    # ---------------------------
    # इ मात्रा reverse (consonant + ि -> f + consonant)
    # ---------------------------
    txt = re.sub(r"([क-ह])ि", lambda m: "f" + m.group(1), txt)

    # ---------------------------
    # MULTI-LATTER (map multi-char sequences first)
    # ---------------------------
    for u, k in base_map.items():
        if len(u) > 1:
            txt = txt.replace(u, k)

    # ---------------------------
    # HALF LETTERS (like क्)
    # ---------------------------
    for u in list(base_map.keys()):
        half = u + "्"
        if half in txt:
            txt = txt.replace(half, base_map[u])

    # ---------------------------
    # SINGLE LETTERS
    # ---------------------------
    for u, k in base_map.items():
        if len(u) == 1:
            txt = txt.replace(u, k)

    # ---------------------------
    # MATRAS
    # ---------------------------
    for u, k in matras.items():
        txt = txt.replace(u, k)

    return txt.replace("्", "")


# ---------------------------------------------------------
#   SPLIT HINDI / ENGLISH SEGMENTS
# ---------------------------------------------------------
def split_segments(text):
    text = normalize_spaces(text)

    output = []
    buf = ""
    hindi = False

    def flush():
        nonlocal buf, hindi
        if buf:
            if hindi:
                output.append(("hindi", convert_hindi(buf)))
            else:
                output.append(("eng", buf))
        buf = ""

    for ch in text:
        if ch == " ":
            flush()
            output.append(("space", " "))
            continue

        if "\u0900" <= ch <= "\u097F":
            if not hindi:
                flush()
                hindi = True
            buf += ch
        else:
            if hindi:
                flush()
                hindi = False
            buf += ch

    flush()
    return output


# ---------------------------------------------------------
#    PUBLIC: convert(text) -> returns final converted string
# ---------------------------------------------------------
def convert(text: str) -> str:
    """
    Main function used by Streamlit app.
    Returns Kruti-encoded string where Hindi segments converted and English kept as-is.
    """
    parts = split_segments(text)
    out = []
    for typ, txt in parts:
        if typ == "hindi":
            out.append(txt)
        elif typ == "eng":
            out.append(txt)
        else:  # space
            out.append(txt)
    return "".join(out)


# ---------------------------------------------------------
#    PUBLIC: build_docx_bytes(text) -> returns .docx bytes
# ---------------------------------------------------------
def build_docx_bytes(text: str) -> bytes:
    """
    Build a DOCX file in-memory using python-docx.
    Hindi segments are set to 'Kruti Dev 010' font name (like original).
    Returns bytes (useful for Streamlit download_button or sending file).
    """
    doc = Document()
    lines = text.split("\n")

    for line in lines:
        p = doc.add_paragraph()
        parts = split_segments(line)
        for typ, txt in parts:
            run = p.add_run(txt)
            run._element.set(qn("xml:space"), "preserve")
            if typ == "hindi":
                try:
                    run.font.name = "Kruti Dev 010"
                    run.font.size = Pt(18)
                    r = run.element.rPr.rFonts
                    r.set(qn("w:eastAsia"), "Kruti Dev 010")
                except Exception:
                    # ignore font-setting errors
                    pass
            elif typ == "eng":
                run.font.name = "Times New Roman"
                run.font.size = Pt(15)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


# Aliases to increase chance Streamlit app finds one of these names
convert_text = convert
to_krutidev = convert
mangal_to_krutidev = convert
unicode_to_krutidev = convert
main = convert


# Quick local test when run directly
if __name__ == "__main__":
    sample = "यहाँ एक उदाहरण वाक्य है। Hello world."
    print(convert(sample))
